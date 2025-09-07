"""
DocParse - Professional Document Processing Library
==================================================

A comprehensive document processing library for extracting text content
from PDF, TXT, and DOCX files with enterprise-grade features.

Author: DocParse Team
Version: 1.0.0
License: MIT
"""

import os
import logging
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, asdict
from enum import Enum

import pdfplumber
import fitz 
from docx import Document
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    TXT = "txt"
    DOCX = "docx"


class ExtractionMethod(Enum):
    """Text extraction methods"""
    PDFPLUMBER = "pdfplumber"
    PYMUPDF = "pymupdf"
    PLAIN_TEXT = "plain_text"
    PYTHON_DOCX = "python_docx"


@dataclass
class ProcessingOptions:
    """Configuration options for document processing"""
    max_pages: Optional[int] = None
    preserve_formatting: bool = False
    extract_tables: bool = True
    extract_metadata: bool = True
    encoding: str = "utf-8"
    error_handling: str = "ignore"  # ignore, strict, replace


@dataclass
class DocumentMetadata:
    """Document metadata container"""
    filename: str
    file_size: int
    total_pages: int
    pages_processed: int
    extraction_timestamp: str
    extraction_method: str
    processing_time: float
    document_type: str
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return asdict(self)


@dataclass
class PageContent:
    """Container for page content"""
    page_number: int
    text: str
    word_count: int
    character_count: int
    tables_count: int = 0
    tables: List[List[List[str]]] = None
    
    def __post_init__(self):
        if self.tables is None:
            self.tables = []
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return asdict(self)


@dataclass
class ExtractionResult:
    """Complete document extraction result"""
    metadata: DocumentMetadata
    pages: List[PageContent]
    success: bool = True
    error_message: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            'metadata': self.metadata.to_dict(),
            'pages': [page.to_dict() for page in self.pages],
            'success': self.success,
            'error_message': self.error_message
        }


class DocumentProcessor(ABC):
    """Abstract base class for document processors"""
    
    @abstractmethod
    def extract_text(self, file_path: Path, options: ProcessingOptions) -> ExtractionResult:
        """Extract text from document"""
        pass
    
    @abstractmethod
    def supports_file_type(self, file_path: Path) -> bool:
        """Check if processor supports this file type"""
        pass


class PDFProcessor(DocumentProcessor):
    """Professional PDF document processor"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.PDFProcessor")
    
    def supports_file_type(self, file_path: Path) -> bool:
        """Check if file is a PDF"""
        return file_path.suffix.lower() == '.pdf'
    
    def extract_text(self, file_path: Path, options: ProcessingOptions) -> ExtractionResult:
        """Extract text from PDF using PDFplumber with PyMuPDF fallback"""
        start_time = datetime.now()
        
        try:
            return self._extract_with_pdfplumber(file_path, options, start_time)
        except Exception as e:
            self.logger.warning(f"PDFplumber failed: {e}. Trying PyMuPDF fallback...")
            try:
                return self._extract_with_pymupdf(file_path, options, start_time)
            except Exception as fallback_error:
                self.logger.error(f"Both extraction methods failed: {e}, {fallback_error}")
                return ExtractionResult(
                    metadata=self._create_error_metadata(file_path, start_time),
                    pages=[],
                    success=False,
                    error_message=f"PDF extraction failed: {str(e)}"
                )
    
    def _extract_with_pdfplumber(self, file_path: Path, options: ProcessingOptions, start_time: datetime) -> ExtractionResult:
        """Extract using PDFplumber"""
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            pages_to_process = min(options.max_pages or total_pages, total_pages)
            
            pages = []
            for page_num in range(pages_to_process):
                page = pdf.pages[page_num]
                text = page.extract_text() or ""
                
                tables = []
                if options.extract_tables:
                    tables = page.extract_tables() or []
                
                page_content = PageContent(
                    page_number=page_num + 1,
                    text=text.strip(),
                    word_count=len(text.split()) if text.strip() else 0,
                    character_count=len(text.strip()),
                    tables_count=len(tables),
                    tables=tables
                )
                pages.append(page_content)
            
            metadata = DocumentMetadata(
                filename=file_path.name,
                file_size=file_path.stat().st_size,
                total_pages=total_pages,
                pages_processed=pages_to_process,
                extraction_timestamp=datetime.now().isoformat(),
                extraction_method=ExtractionMethod.PDFPLUMBER.value,
                processing_time=(datetime.now() - start_time).total_seconds(),
                document_type=DocumentType.PDF.value
            )
            
            return ExtractionResult(metadata=metadata, pages=pages)
    
    def _extract_with_pymupdf(self, file_path: Path, options: ProcessingOptions, start_time: datetime) -> ExtractionResult:
        """Fallback extraction using PyMuPDF"""
        doc = fitz.open(file_path)
        total_pages = len(doc)
        pages_to_process = min(options.max_pages or total_pages, total_pages)
        
        pages = []
        for page_num in range(pages_to_process):
            page = doc[page_num]
            text = page.get_text()
            
            page_content = PageContent(
                page_number=page_num + 1,
                text=text.strip(),
                word_count=len(text.split()) if text.strip() else 0,
                character_count=len(text.strip()),
                tables_count=0,
                tables=[]
            )
            pages.append(page_content)
        
        doc.close()
        
        metadata = DocumentMetadata(
            filename=file_path.name,
            file_size=file_path.stat().st_size,
            total_pages=total_pages,
            pages_processed=pages_to_process,
            extraction_timestamp=datetime.now().isoformat(),
            extraction_method=ExtractionMethod.PYMUPDF.value,
            processing_time=(datetime.now() - start_time).total_seconds(),
            document_type=DocumentType.PDF.value
        )
        
        return ExtractionResult(metadata=metadata, pages=pages)
    
    def _create_error_metadata(self, file_path: Path, start_time: datetime) -> DocumentMetadata:
        """Create metadata for failed extraction"""
        return DocumentMetadata(
            filename=file_path.name,
            file_size=file_path.stat().st_size if file_path.exists() else 0,
            total_pages=0,
            pages_processed=0,
            extraction_timestamp=datetime.now().isoformat(),
            extraction_method="failed",
            processing_time=(datetime.now() - start_time).total_seconds(),
            document_type=DocumentType.PDF.value
        )


class TXTProcessor(DocumentProcessor):
    """Professional TXT document processor"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.TXTProcessor")
    
    def supports_file_type(self, file_path: Path) -> bool:
        """Check if file is a TXT"""
        return file_path.suffix.lower() == '.txt'
    
    def extract_text(self, file_path: Path, options: ProcessingOptions) -> ExtractionResult:
        """Extract text from TXT file"""
        start_time = datetime.now()
        
        try:
            with open(file_path, 'r', encoding=options.encoding, errors=options.error_handling) as file:
                content = file.read()
            
            # Smart page splitting
            pages_text = self._split_into_pages(content)
            total_pages = len(pages_text)
            pages_to_process = min(options.max_pages or total_pages, total_pages)
            
            pages = []
            for page_num in range(pages_to_process):
                text = pages_text[page_num].strip()
                
                page_content = PageContent(
                    page_number=page_num + 1,
                    text=text,
                    word_count=len(text.split()) if text else 0,
                    character_count=len(text),
                    tables_count=0,
                    tables=[]
                )
                pages.append(page_content)
            
            metadata = DocumentMetadata(
                filename=file_path.name,
                file_size=file_path.stat().st_size,
                total_pages=total_pages,
                pages_processed=pages_to_process,
                extraction_timestamp=datetime.now().isoformat(),
                extraction_method=ExtractionMethod.PLAIN_TEXT.value,
                processing_time=(datetime.now() - start_time).total_seconds(),
                document_type=DocumentType.TXT.value
            )
            
            return ExtractionResult(metadata=metadata, pages=pages)
            
        except Exception as e:
            self.logger.error(f"TXT extraction failed: {e}")
            return ExtractionResult(
                metadata=self._create_error_metadata(file_path, start_time),
                pages=[],
                success=False,
                error_message=f"TXT extraction failed: {str(e)}"
            )
    
    def _split_into_pages(self, content: str) -> List[str]:
        """Split text content into logical pages"""
        # Try form feed characters first
        pages = content.split('\f')
        if len(pages) > 1:
            return [page for page in pages if page.strip()]
        
        # Try triple newlines
        pages = content.split('\n\n\n')
        if len(pages) > 1:
            return [page for page in pages if page.strip()]
        
        # Split by character count (approximately 3000 chars per page)
        if len(content) > 3000:
            pages = []
            for i in range(0, len(content), 3000):
                pages.append(content[i:i+3000])
            return pages
        
        # Return as single page
        return [content] if content.strip() else ['']
    
    def _create_error_metadata(self, file_path: Path, start_time: datetime) -> DocumentMetadata:
        """Create metadata for failed extraction"""
        return DocumentMetadata(
            filename=file_path.name,
            file_size=file_path.stat().st_size if file_path.exists() else 0,
            total_pages=0,
            pages_processed=0,
            extraction_timestamp=datetime.now().isoformat(),
            extraction_method="failed",
            processing_time=(datetime.now() - start_time).total_seconds(),
            document_type=DocumentType.TXT.value
        )


class DOCXProcessor(DocumentProcessor):
    """Professional DOCX document processor"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DOCXProcessor")
    
    def supports_file_type(self, file_path: Path) -> bool:
        """Check if file is a DOCX"""
        return file_path.suffix.lower() == '.docx'
    
    def extract_text(self, file_path: Path, options: ProcessingOptions) -> ExtractionResult:
        """Extract text from DOCX file"""
        start_time = datetime.now()
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs and group into pages
            pages_text = self._group_into_pages(doc, options)
            total_pages = len(pages_text)
            pages_to_process = min(options.max_pages or total_pages, total_pages)
            
            pages = []
            for page_num in range(pages_to_process):
                text = pages_text[page_num]
                
                page_content = PageContent(
                    page_number=page_num + 1,
                    text=text,
                    word_count=len(text.split()) if text else 0,
                    character_count=len(text),
                    tables_count=0,
                    tables=[]
                )
                pages.append(page_content)
            
            metadata = DocumentMetadata(
                filename=file_path.name,
                file_size=file_path.stat().st_size,
                total_pages=total_pages,
                pages_processed=pages_to_process,
                extraction_timestamp=datetime.now().isoformat(),
                extraction_method=ExtractionMethod.PYTHON_DOCX.value,
                processing_time=(datetime.now() - start_time).total_seconds(),
                document_type=DocumentType.DOCX.value
            )
            
            return ExtractionResult(metadata=metadata, pages=pages)
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return ExtractionResult(
                metadata=self._create_error_metadata(file_path, start_time),
                pages=[],
                success=False,
                error_message=f"DOCX extraction failed: {str(e)}"
            )
    
    def _group_into_pages(self, doc: Document, options: ProcessingOptions) -> List[str]:
        """Group paragraphs into logical pages"""
        pages = []
        current_page = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                current_page.append(text)
            elif current_page:  # Empty paragraph might indicate page break
                pages.append('\n'.join(current_page))
                current_page = []
        
        # Add remaining content
        if current_page:
            pages.append('\n'.join(current_page))
        
        # If no natural breaks, group by paragraph count
        if not pages and doc.paragraphs:
            all_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            pages = [all_text] if all_text else ['']
        
        return pages or ['']
    
    def _create_error_metadata(self, file_path: Path, start_time: datetime) -> DocumentMetadata:
        """Create metadata for failed extraction"""
        return DocumentMetadata(
            filename=file_path.name,
            file_size=file_path.stat().st_size if file_path.exists() else 0,
            total_pages=0,
            pages_processed=0,
            extraction_timestamp=datetime.now().isoformat(),
            extraction_method="failed",
            processing_time=(datetime.now() - start_time).total_seconds(),
            document_type=DocumentType.DOCX.value
        )


class DocParseEngine:
    """Professional document parsing engine"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocParseEngine")
        self.processors = {
            DocumentType.PDF: PDFProcessor(),
            DocumentType.TXT: TXTProcessor(),
            DocumentType.DOCX: DOCXProcessor()
        }
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return [doc_type.value for doc_type in DocumentType]
    
    def is_supported_file(self, file_path: Union[str, Path]) -> bool:
        """Check if file type is supported"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower().lstrip('.')
        
        try:
            doc_type = DocumentType(extension)
            return self.processors[doc_type].supports_file_type(file_path)
        except ValueError:
            return False
    
    def extract_text(self, file_path: Union[str, Path], options: Optional[ProcessingOptions] = None) -> ExtractionResult:
        """Extract text from supported document"""
        file_path = Path(file_path)
        options = options or ProcessingOptions()
        
        self.logger.info(f"Starting extraction for: {file_path.name}")
        
        # Validate file exists
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            return ExtractionResult(
                metadata=DocumentMetadata(
                    filename=file_path.name,
                    file_size=0,
                    total_pages=0,
                    pages_processed=0,
                    extraction_timestamp=datetime.now().isoformat(),
                    extraction_method="failed",
                    processing_time=0.0,
                    document_type="unknown"
                ),
                pages=[],
                success=False,
                error_message="File not found"
            )
        
        # Get document type
        extension = file_path.suffix.lower().lstrip('.')
        try:
            doc_type = DocumentType(extension)
        except ValueError:
            self.logger.error(f"Unsupported file type: {extension}")
            return ExtractionResult(
                metadata=DocumentMetadata(
                    filename=file_path.name,
                    file_size=file_path.stat().st_size,
                    total_pages=0,
                    pages_processed=0,
                    extraction_timestamp=datetime.now().isoformat(),
                    extraction_method="failed",
                    processing_time=0.0,
                    document_type=extension
                ),
                pages=[],
                success=False,
                error_message=f"Unsupported file type: {extension}"
            )
        
        # Process document
        processor = self.processors[doc_type]
        result = processor.extract_text(file_path, options)
        
        if result.success:
            self.logger.info(f"Successfully extracted {result.metadata.pages_processed} pages from {file_path.name}")
        else:
            self.logger.error(f"Extraction failed for {file_path.name}: {result.error_message}")
        
        return result


# Convenience functions for easy usage
def extract_document(file_path: Union[str, Path], max_pages: Optional[int] = None, **kwargs) -> ExtractionResult:
    """Convenience function to extract text from a document"""
    options = ProcessingOptions(max_pages=max_pages, **kwargs)
    engine = DocParseEngine()
    return engine.extract_text(file_path, options)


def get_supported_formats() -> List[str]:
    """Get list of supported document formats"""
    return DocParseEngine().get_supported_extensions()


# Version info
__version__ = "1.0.0"
__author__ = "DocParse Team"
__license__ = "MIT"
